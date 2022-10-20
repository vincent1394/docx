# This is a sample Python script.

# Press ⌃R to execute it or replace it with your code.
# Press Double ⇧ to search everywhere for classes, files, tool windows, actions, and settings.
import json
#refence  https://www.codenong.com/1665454009794833226/
#refence  https://python-docx.readthedocs.io/en/latest/user/sections.html
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx.shared import Cm, Pt


def print_hi(name):
    # Use a breakpoint in the code line below to debug your script.
    print(f'Hi, {name}')  # Press ⌘F8 to toggle the breakpoint.


# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    filename = '/Users/bjhl/Downloads/测试test.docx'
    #f = open(filename, 'rb')
    doc = Document(filename)
    #doc.settings.odd_and_even_pages_header_footer= True
    idx = 0
    for section in doc.sections:
        section.orientation = WD_ORIENTATION.PORTRAIT
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.clear()
        #paragraph.style= WD_PARAGRAPH_ALIGNMENT.CENTER
        #run = paragraph.add_run("\n")
        section.top_margin = Cm(4.5)
        section.header_distance = Cm(2)
        run = paragraph.add_run()
        title = "    "
        if idx %2 == 0:
            section.left_margin=Cm(2.5)
            # paragraph.style = WD_STYLE_TYPE.PARAGRAPH
            run.add_picture("1666157071516.jpg", Cm(5.28), Cm(1.22))
            run = paragraph.add_run("\n")
            run.add_picture("1666157086632.jpg", Cm(1), Cm(0.2))
            if 0 <=idx<23:
                title +="第一章 标本采集与制作"
            elif 23 <= idx<42:
                title += "第二章	数据获取与分析"
            elif 42 <= idx<81:
                title += "第三章	实验方案与流程"
            else:
                title += "第四章	观点评述与综述"
            run = paragraph.add_run("\t")
            run.font.size = Pt(12)
            run.bold = True  # 字体是否加粗
            run.font.name = 'Times New Roman'
            run.alignment = "JUSTIFY"
            run.add_text(title)
        else:
            # paragraph.style = WD_STYLE_TYPE.PARAGRAPH
            run.add_picture("1666157071516.jpg", Cm(3.4), Cm(0.79))#LOGO
            section.left_margin = Cm(3.8)
            run = paragraph.add_run("\n")
            run.add_picture("1666157086632.jpg", Cm(1), Cm(0.2)) #LOGO UNDERLINE

            title += "Zoological Systematics and Evolution Protocol eBook"
            run = paragraph.add_run("\t")
            run.font.size = Pt(10)
            run.bold = True  # 字体是否加粗
            run.font.name = 'Times New Roman'
            run.alignment = "JUSTIFY"
            run.add_text(title)
            run = paragraph.add_run("\n")
            run.add_picture("1666168043950.jpg",Cm(15.6), Cm(0.05))#SOID UNDERLINE
        idx += 1
    doc.save("8.docx")

# See PyCharm help at https://www.jetbrains.com/help/pycharm/
